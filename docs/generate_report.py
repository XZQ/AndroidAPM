"""Generate the tracked AndroidAPM DOCX reports from repository-backed facts."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOCS_DIR = Path(__file__).resolve().parent
DIAGRAM_DIR = DOCS_DIR / "architecture" / "generated-diagrams"
REPORT_DATE = "2026-07-23"
RUNTIME_BASELINE = "develop（以 git log 为准）"


def set_cell_shading(cell, fill: str) -> None:
    """Apply a hexadecimal background color to a table cell."""
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, value: str, *, bold: bool = False, color: str | None = None) -> None:
    """Replace a cell's text while keeping report typography consistent."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.size = Pt(9)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_document(document: Document, title: str, subtitle: str) -> None:
    """Configure page layout, styles, and a compact title block."""
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for style_name, size, color in (
        ("Title", 24, "12355B"),
        ("Heading 1", 16, "12355B"),
        ("Heading 2", 12, "1F5A94"),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    heading = document.add_paragraph(style="Title")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run(title)
    subheading = document.add_paragraph()
    subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subheading.add_run(subtitle)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("52657A")
    document.add_paragraph()


def add_summary_table(document: Document) -> None:
    """Add the current source inventory and platform baseline."""
    rows = [
        ("构建单元", "27", "25 个根子项目 + apm-plugin + build-logic"),
        ("主源码", "164", "159 Kotlin + 4 C + 1 proto"),
        ("测试文件", "102", "JVM、Robolectric、instrumented benchmark、device-soak host gate、插件和 native 契约测试"),
        ("Android", "compile 34 / min 24", "targetSdk 34"),
        ("构建栈", "Java 17 toolchain / Gradle 8.13", "Gradle runtime JDK 17+ / AGP 8.13.2 / Kotlin 2.2.21"),
        ("运行时代码基线", RUNTIME_BASELINE, "以当前源码和可执行验证为准"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for index, value in enumerate(("项目", "当前值", "说明")):
        set_cell_text(table.rows[0].cells[index], value, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[index], "12355B")
    for label, value, note in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], value)
        set_cell_text(cells[2], note)


def add_bullets(document: Document, items: list[str]) -> None:
    """Add compact bullet paragraphs."""
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_capability_table(document: Document) -> None:
    """Describe capabilities without presenting configured flags as implementations."""
    rows = [
        ("生产安全入口", "Strict profile、显式 consent、撤回清理", "初始化前 fail closed；停止 delivery 后清理 queue/outbox/IPC"),
        ("Collector Wire V2", "Typed fields、standard resource、batch ID、size、exact ACK", "Legacy wire 不变；V2 独立 schema 与 endpoint"),
        ("关键事件与损失证据", "Crash/ANR sync hand-off、drop reason/priority", "绕过共享队列且不做网络；未知 priority 显式 unattributed"),
        ("自动生命周期接入", "Memory、Crash、ANR、Launch、FPS、GC、Render、Thread", "SDK 初始化后可运行；仍受权限、API 和设备限制"),
        ("时间与快照语义", "epoch collector 时间 + 单调 duration/window；异步事件 map 冻结", "避免系统时间跳变和宿主后续修改污染已发生事件"),
        ("跨层字节预算", "Dispatcher 8 MiB；IPC 4 MiB/256 KiB/1 MiB/16 MiB；SQLite 256 KiB/64 MiB", "各层按 retained estimate、encoded/file bytes、durable payload 的真实资源维度独立限界"),
        ("真机开销门", "A/B 启动、主线程、CPU、PSS、功耗、磁盘、热、离线重启", "三项 microbenchmark 通过；Redmi 修复后两轮原预算 smoke 通过；OnePlus Android 16 以限定包重装回退通过并产出 29.062 mAh/hour UID 功耗；24h/72h 未执行"),
        ("显式 API 接入", "Network、SQLite、IPC、WebView、ThreadPool、Battery、IO", "由宿主在真实调用点安装 wrapper 或传入 executor/耗时/错误"),
        ("构建期插桩", "ASM slow-method", "AGP instrumentation API；需应用 Gradle 插件"),
        ("事件管线", "eventId → Dispatcher → SQLite claim lease → Uploader", "owner 确认成功后删除，语义为至少一次"),
        ("扩展", "Trace、OTel exporter", "Trace 为进程内 span；OTel exporter 仅做事件映射"),
        ("不支持的全局 Hook", "Binder hidden Hook、WebView 全局接管、通用线程泄漏、GPU overdraw", "兼容字段已 deprecated/false；不伪装成自动能力"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for index, value in enumerate(("接入形态", "范围", "真实边界")):
        set_cell_text(table.rows[0].cells[index], value, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[index], "1F5A94")
    for kind, scope, boundary in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], kind, bold=True)
        set_cell_text(cells[1], scope)
        set_cell_text(cells[2], boundary)


def add_diagram(document: Document, filename: str, caption: str) -> None:
    """Embed a generated diagram when the PNG exists."""
    path = DIAGRAM_DIR / filename
    if not path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(6.7))
    caption_paragraph = document.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption_paragraph.runs:
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("52657A")


def add_footer(document: Document) -> None:
    """Add a source-of-truth footer to each section."""
    for section in document.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(
            f"Generated {REPORT_DATE} from the local AndroidAPM repository; "
            "volatile status belongs in docs/Android_APM_项目文档.md"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("6B7785")


def build_status_report() -> Document:
    """Build a source-backed current-state and gap report."""
    document = Document()
    configure_document(
        document,
        "AndroidAPM 当前能力与生产差距报告",
        f"本地源码事实快照 · {REPORT_DATE}",
    )

    document.add_heading("结论先行", level=1)
    document.add_paragraph(
        "AndroidAPM 已形成可构建、可测试的多模块 Android 客户端 APM SDK：采集事件进入有界异步管线，"
        "通过 SQLite outbox 持久化并交给可注入上传器。它仍是客户端框架，不包含生产采集后端、查询、"
        "告警和运营闭环。2026-07-23 的物理设备 microbenchmark 通过；FPS 静态页面 VSync observer 修复后，"
        "两轮 smoke 在原 20% 门禁下以 12.928% / 12.362% CPU 通过；OnePlus Android 16 预检也在原预算下通过并产出 UID 功耗证据。"
        "短 smoke 不替代 24h/72h 长稳验收。"
    )
    add_summary_table(document)

    document.add_heading("能力矩阵", level=1)
    add_capability_table(document)

    document.add_heading("数据交付语义", level=1)
    add_bullets(
        document,
        [
            "Strict production profile 在创建 SDK 资源前校验显式 consent、HTTPS/自定义 uploader、SQLite、PII 与 debug logging。",
            "Consent revoke 不做优雅 drain；停止 delivery 后清理活动队列、SQLite/File outbox 与 IPC hand-off 文件。",
            "V2 protobuf envelope 按实际编码字节拆批；2xx 只有 schema、batchId 和事件数精确 ACK 才删除 outbox。",
            "Dispatcher 同时受 2048 条和 8 MiB 估算字节预算约束；IPC 与 SQLite 另按 encoded/file/durable bytes 独立限界。",
            "SQLite outbox 支持进程重启后的持久化恢复，并在上传确认成功后删除。",
            "当前语义是 acknowledged at-least-once；稳定 eventId 已贯穿 wire/storage，服务端仍须幂等。",
            "SQLite 写事务提供多消费者 claim/lease/expiry、owner-aware ACK/failure 和 shutdown release。",
            "持久化 codec v3 为常用标量写类型标签并兼容读取 v1/v2；任意对象仍安全降级为字符串。",
            "apm-uploader 保持底层依赖方向，同时以模块内命名工厂显式治理 worker/scheduler 的 daemon 与 MIN_PRIORITY。",
            "Dispatcher 在 self-monitor 开启时以固定无逐样本分配直方图输出六阶段 count、平均、P95 上界和最大延迟；"
            "关闭时跳过计时，阶段证据不自动触发并行化。",
        ],
    )
    add_diagram(document, "android-apm-event-pipeline.png", "图 1：事件从采集到确认删除的真实管线")

    document.add_heading("走向生产的优先级", level=1)
    priorities = [
        ("P1", "在原预算 smoke 已连续通过后，执行 24h/72h、长稳功耗、热与磁盘验收"),
        ("P0", "接入生产 collector，并定义鉴权、限流、协议版本和隐私治理"),
        ("P0", "在 Collector 按客户端 eventId 幂等，明确整批 ACK、重放与死信"),
        ("P1", "建立真机/OEM/API 设备矩阵，覆盖 native、ANR、多进程和长期离线"),
        ("P1", "建设 Native 符号上传/后台符号化与外部制品发布"),
        ("P2", "建设查询、聚合、告警、版本对比与 SDK 自身健康观测"),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "优先级", bold=True, color="FFFFFF")
    set_cell_text(table.rows[0].cells[1], "工作项", bold=True, color="FFFFFF")
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "12355B")
    for priority, item in priorities:
        cells = table.add_row().cells
        set_cell_text(cells[0], priority, bold=True)
        set_cell_text(cells[1], item)

    add_footer(document)
    return document


def build_architecture_report() -> Document:
    """Build an implementation-oriented architecture report."""
    document = Document()
    configure_document(
        document,
        "AndroidAPM 架构与模块报告",
        f"代码、构建与集成边界 · {REPORT_DATE}",
    )

    document.add_heading("架构概览", level=1)
    document.add_paragraph(
        "架构以 apm-core 为编排中心，以 apm-model、apm-storage、apm-uploader 形成数据底座；"
        "15 个监控模块负责采集或接收宿主埋点，apm-trace 与 apm-otel-exporter 提供扩展能力，"
        "apm-bundle 提供单依赖完整客户端分发，"
        "apm-plugin 在构建期完成慢方法插桩，apm-benchmark 提供非发布真机开销入口。"
    )
    add_diagram(document, "android-apm-overview.png", "图 1：客户端 SDK 总体结构")

    document.add_heading("模块依赖原则", level=1)
    add_bullets(
        document,
        [
            "监控模块通过 apm-core 上报，不直接操作 SQLite 或 HTTP。",
            "apm-uploader 不反向依赖 apm-core；因此保留模块内执行器和注入式 UploaderLogger，"
            "其 worker/scheduler 由本地命名工厂显式设置 daemon 与 MIN_PRIORITY。",
            "apm-plugin 与 build-logic 是 included build，不属于根 Gradle 子项目。",
            "apm-bundle 只聚合发布依赖，不承载运行时实现，也不自动应用慢方法插件。",
            "sample app 是接入示例和冒烟入口，不是生产 collector。",
            "单 dispatcher worker 保留顺序语义；sdk_health 的六阶段 count/平均/P95 上界/最大延迟字段"
            "用于先归因再决定是否分区或并行。",
            "apm-benchmark 不进入 Maven publication；microbenchmark 固定 time/allocation，device-soak 固定 A/B/资源/时长/重启证据，并区分 enabled 绝对 CPU 门禁与 control/delta 归因；OEM 禁止 pm clear 时只对所选 sample APK 卸载重装并记录 provenance。",
            "device-soak 仅对只读证据命令做三类 ADB transport 瞬断的有界重试；安装、卸载、清理和 Activity 命令不自动重放，工件保留 retry count。",
        ],
    )
    add_diagram(document, "android-apm-module-dependencies.png", "图 2：主要模块依赖方向")

    document.add_heading("监控模块与接入方式", level=1)
    add_capability_table(document)
    add_diagram(document, "android-apm-monitoring-modules.png", "图 3：监控模块按接入方式分组")

    document.add_heading("慢方法插桩", level=1)
    document.add_paragraph(
        "apm-plugin 使用 AGP instrumentation API 访问字节码，在命中的方法入口和出口写入计时逻辑。"
        "是否插桩由插件配置、包过滤和阈值共同控制；该能力需要宿主显式应用插件，不能描述为全局零侵入。"
    )
    add_diagram(document, "android-apm-slow-method-instrumentation.png", "图 4：构建期插桩与运行时上报")

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("验证入口", level=1)
    document.add_paragraph("仓库的标准验证链如下，最终结果以项目状态文档和实际命令输出为准：")
    for command in (
        "./gradlew assembleDebug",
        "./gradlew testDebugUnitTest",
        "./gradlew -p apm-plugin test",
        "./gradlew :apm-benchmark:assembleRelease :apm-benchmark:compileReleaseAndroidTestKotlin",
        "./gradlew :apm-benchmark:verifyReleasePerformanceBudgets",
        "python -m unittest discover -s apm-benchmark/tests -p test_*.py",
        "./gradlew :apm-benchmark:verifyDeviceSoakFromResults -PapmDeviceSoakResults=<json> -PapmDeviceSoakProfile=<profile>",
        "./gradlew lintDebug",
        "./gradlew assembleRelease",
        "./gradlew publishToMavenLocal",
        "./gradlew -p smoke-tests/maven-consumer clean assembleDebug",
    ):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(command)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    document.add_heading("文档治理", level=1)
    add_bullets(
        document,
        [
            "当前状态唯一主入口：docs/Android_APM_项目文档.md。",
            "便携交接入口：docs/PROJECT_HANDOFF.md。",
            "架构细节：docs/architecture/00_整体架构.md 与对应模块文档。",
            "云端、发布和真机设备实验室清单由独立 AndroidAPM-Server 仓库的 docs/云端待建设清单.md 维护。",
            "DOCX、SVG 和 PNG 是派生产物；出现冲突时以源码和 Markdown 为准。",
            "历史原始附件已移除，不再作为文档交付或当前事实来源。",
        ],
    )

    add_footer(document)
    return document


def save(document: Document, filename: str) -> Path:
    """Save a report beside this generator and return its path."""
    output = DOCS_DIR / filename
    document.save(output)
    return output


def main() -> None:
    """Generate both tracked reports."""
    outputs = [
        save(build_status_report(), "APM_对比报告.docx"),
        save(build_architecture_report(), "APM_框架对比报告.docx"),
    ]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
